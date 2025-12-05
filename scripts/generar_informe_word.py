#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script para generar documento Word profesional del Informe Técnico de Justificación
de Reducción de Cantidades ITS - Proyecto APP Puerto Salgar - Barrancabermeja
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import re

# Colores corporativos
COLOR_AZUL = RGBColor(5, 99, 193)  # #0563C1
COLOR_GRIS = RGBColor(68, 84, 106)  # #44546A
COLOR_AMARILLO = RGBColor(255, 242, 204)  # #FFF2CC
COLOR_VERDE = RGBColor(226, 239, 218)  # #E2EFDA
COLOR_BORDE_AMARILLO = RGBColor(191, 143, 0)  # #BF8F00
COLOR_BORDE_VERDE = RGBColor(112, 173, 71)  # #70AD47
COLOR_BORDE_GRIS = RGBColor(191, 191, 191)  # #BFBFBF
COLOR_FONDO_TABLA = RGBColor(242, 242, 242)  # #F2F2F2
COLOR_CITA = RGBColor(231, 230, 230)  # #E7E6E6

def crear_estilos(doc):
    """Crea estilos personalizados para el documento"""
    styles = doc.styles
    
    # Estilo Título 1
    try:
        style_h1 = styles['Heading 1']
    except:
        style_h1 = styles.add_style('Custom Heading 1', WD_STYLE_TYPE.PARAGRAPH)
    style_h1.font.name = 'Calibri'
    style_h1.font.size = Pt(16)
    style_h1.font.bold = True
    style_h1.font.color.rgb = COLOR_AZUL
    style_h1.font.all_caps = True
    
    # Estilo Título 2
    try:
        style_h2 = styles['Heading 2']
    except:
        style_h2 = styles.add_style('Custom Heading 2', WD_STYLE_TYPE.PARAGRAPH)
    style_h2.font.name = 'Calibri'
    style_h2.font.size = Pt(14)
    style_h2.font.bold = True
    style_h2.font.color.rgb = COLOR_AZUL
    
    # Estilo Título 3
    try:
        style_h3 = styles['Heading 3']
    except:
        style_h3 = styles.add_style('Custom Heading 3', WD_STYLE_TYPE.PARAGRAPH)
    style_h3.font.name = 'Calibri'
    style_h3.font.size = Pt(12)
    style_h3.font.bold = True
    style_h3.font.color.rgb = COLOR_GRIS
    
    # Estilo Título 4
    try:
        style_h4 = styles['Heading 4']
    except:
        style_h4 = styles.add_style('Custom Heading 4', WD_STYLE_TYPE.PARAGRAPH)
    style_h4.font.name = 'Calibri'
    style_h4.font.size = Pt(11)
    style_h4.font.bold = True
    style_h4.font.italic = True
    style_h4.font.color.rgb = COLOR_GRIS
    
    # Estilo Normal
    style_normal = styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(11)
    style_normal.paragraph_format.line_spacing = 1.15
    style_normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def agregar_portada(doc):
    """Agrega portada profesional al documento"""
    # Título principal
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('INFORME TÉCNICO DE JUSTIFICACIÓN\nREDUCCIÓN DE CANTIDADES ITS')
    run.font.name = 'Calibri'
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = COLOR_AZUL
    
    doc.add_paragraph()  # Espacio
    
    # Subtítulo
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('PROYECTO APP PUERTO SALGAR - BARRANCABERMEJA')
    run.font.name = 'Calibri'
    run.font.size = Pt(14)
    run.font.bold = True
    
    # Espacios
    for _ in range(3):
        doc.add_paragraph()
    
    # Información del documento
    info = [
        ('Versión:', '1.0'),
        ('Fecha:', datetime.now().strftime('%d de %B de %Y')),
        ('Estado:', 'Completado'),
    ]
    
    for label, value in info:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'{label} {value}')
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
    
    # Espacios
    for _ in range(3):
        doc.add_paragraph()
    
    # Elaborado por
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Elaborado por:\nAdministrador Contractual EPC')
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    
    # Salto de página
    doc.add_page_break()

def agregar_tabla_con_formato(doc, datos, encabezados, titulo=None):
    """Crea una tabla con formato profesional"""
    if titulo:
        p = doc.add_paragraph()
        run = p.add_run(f'Tabla {titulo}')
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.italic = True
        doc.add_paragraph()
    
    table = doc.add_table(rows=1, cols=len(encabezados))
    table.style = 'Light Grid Accent 1'
    
    # Encabezados
    header_cells = table.rows[0].cells
    for i, encabezado in enumerate(encabezados):
        header_cells[i].text = encabezado
        header_cells[i].paragraphs[0].runs[0].font.bold = True
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        # Fondo azul para encabezados
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), '0563C1')
        header_cells[i]._element.get_or_add_tcPr().append(shading_elm)
        header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_cells[i].vertical_alignment = 1  # Centrado vertical
    
    # Datos
    for row_idx, fila in enumerate(datos):
        row_cells = table.add_row().cells
        for col_idx, valor in enumerate(fila):
            row_cells[col_idx].text = str(valor)
            # Filas alternas
            if row_idx % 2 == 1:
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'F2F2F2')
                row_cells[col_idx]._element.get_or_add_tcPr().append(shading_elm)
            # Alineación: números a la derecha, texto a la izquierda
            if col_idx > 0 and any(c.isdigit() for c in str(valor)):
                row_cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                row_cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    doc.add_paragraph()  # Espacio después de la tabla
    return table

def agregar_cita_contractual(doc, texto):
    """Agrega una cita contractual con formato especial"""
    p = doc.add_paragraph()
    p.style = 'Quote'
    run = p.add_run(texto)
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.italic = True
    
    # Fondo gris y borde azul izquierdo
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'E7E6E6')
    p._element.get_or_add_pPr().append(shading_elm)
    
    # Borde izquierdo azul
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pBdr_left = OxmlElement('w:left')
    pBdr_left.set(qn('w:val'), 'single')
    pBdr_left.set(qn('w:sz'), '24')  # 3pt
    pBdr_left.set(qn('w:space'), '0')
    pBdr_left.set(qn('w:color'), '0563C1')
    pBdr.append(pBdr_left)
    pPr.append(pBdr)
    
    doc.add_paragraph()  # Espacio

def agregar_nota_importante(doc, texto):
    """Agrega una nota importante con fondo amarillo"""
    p = doc.add_paragraph()
    run = p.add_run(f'NOTA: {texto}')
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.bold = True
    
    # Fondo amarillo y borde
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'FFF2CC')
    p._element.get_or_add_pPr().append(shading_elm)
    
    # Borde
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side in ['top', 'bottom', 'left', 'right']:
        side_elm = OxmlElement(f'w:{side}')
        side_elm.set(qn('w:val'), 'single')
        side_elm.set(qn('w:sz'), '4')  # 0.5pt
        side_elm.set(qn('w:space'), '0')
        side_elm.set(qn('w:color'), 'BF8F00')
        pBdr.append(side_elm)
    pPr.append(pBdr)
    
    doc.add_paragraph()  # Espacio

def agregar_conclusion(doc, texto):
    """Agrega una conclusión con fondo verde"""
    p = doc.add_paragraph()
    run = p.add_run(f'CONCLUSIÓN: {texto}')
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.bold = True
    
    # Fondo verde y borde
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'E2EFDA')
    p._element.get_or_add_pPr().append(shading_elm)
    
    # Borde
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side in ['top', 'bottom', 'left', 'right']:
        side_elm = OxmlElement(f'w:{side}')
        side_elm.set(qn('w:val'), 'single')
        side_elm.set(qn('w:sz'), '4')
        side_elm.set(qn('w:space'), '0')
        side_elm.set(qn('w:color'), '70AD47')
        pBdr.append(side_elm)
    pPr.append(pBdr)
    
    doc.add_paragraph()  # Espacio

def procesar_markdown_a_word(doc, contenido_md):
    """Procesa el contenido Markdown y lo convierte a Word"""
    lineas = contenido_md.split('\n')
    i = 0
    
    while i < len(lineas):
        linea = lineas[i].strip()
        
        # Saltar líneas vacías
        if not linea:
            i += 1
            continue
        
        # Título 1 (##)
        if linea.startswith('## ') and not linea.startswith('###'):
            texto = linea[3:].strip()
            p = doc.add_heading(texto, level=1)
            p.style.font.name = 'Calibri'
            p.style.font.size = Pt(16)
            p.style.font.bold = True
            p.style.font.color.rgb = COLOR_AZUL
            p.style.font.all_caps = True
        
        # Título 2 (###)
        elif linea.startswith('### '):
            texto = linea[4:].strip()
            p = doc.add_heading(texto, level=2)
            p.style.font.name = 'Calibri'
            p.style.font.size = Pt(14)
            p.style.font.bold = True
            p.style.font.color.rgb = COLOR_AZUL
        
        # Título 3 (####)
        elif linea.startswith('#### '):
            texto = linea[5:].strip()
            p = doc.add_heading(texto, level=3)
            p.style.font.name = 'Calibri'
            p.style.font.size = Pt(12)
            p.style.font.bold = True
            p.style.font.color.rgb = COLOR_GRIS
        
        # Tabla Markdown
        elif linea.startswith('|') and '|' in linea:
            # Leer tabla completa
            tabla_lineas = []
            while i < len(lineas) and lineas[i].strip().startswith('|'):
                tabla_lineas.append(lineas[i].strip())
                i += 1
            i -= 1  # Retroceder una línea
            
            # Procesar tabla
            if len(tabla_lineas) >= 2:
                encabezados = [cell.strip() for cell in tabla_lineas[0].split('|')[1:-1]]
                datos = []
                for fila in tabla_lineas[2:]:  # Saltar separador
                    if '---' not in fila:
                        datos.append([cell.strip() for cell in fila.split('|')[1:-1]])
                
                if datos:
                    agregar_tabla_con_formato(doc, datos, encabezados)
        
        # Cita (bloque >)
        elif linea.startswith('>'):
            texto_cita = linea[1:].strip()
            # Continuar leyendo si hay más líneas de cita
            while i + 1 < len(lineas) and lineas[i + 1].strip().startswith('>'):
                i += 1
                texto_cita += '\n' + lineas[i].strip()[1:].strip()
            agregar_cita_contractual(doc, texto_cita)
        
        # Lista
        elif linea.startswith('- ') or linea.startswith('* '):
            texto = linea[2:].strip()
            p = doc.add_paragraph(texto, style='List Bullet')
            p.style.font.name = 'Calibri'
            p.style.font.size = Pt(11)
        
        # Párrafo normal
        else:
            # Limpiar emojis y símbolos especiales
            texto = linea
            texto = texto.replace('✅', '✓')
            texto = texto.replace('⚠️', '△')
            texto = texto.replace('📄', '')
            texto = texto.replace('⚖️', '')
            texto = texto.replace('🧭', '')
            texto = texto.replace('**', '')  # Negrita se maneja diferente
            
            p = doc.add_paragraph(texto)
            p.style.font.name = 'Calibri'
            p.style.font.size = Pt(11)
            p.style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        i += 1

def agregar_numeracion_paginas(doc):
    """Agrega numeración de páginas al pie de página"""
    section = doc.sections[0]
    footer = section.footer
    
    # Crear párrafo para numeración
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Agregar campo de número de página
    run = p.add_run('Página ')
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    
    # Campo PAGE
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    p._element.append(fldChar1)
    p._element.append(instrText)
    p._element.append(fldChar2)
    
    run = p.add_run(' de ')
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    
    # Campo NUMPAGES
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'begin')
    
    instrText2 = OxmlElement('w:instrText')
    instrText2.set(qn('xml:space'), 'preserve')
    instrText2.text = 'NUMPAGES'
    
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
    
    p._element.append(fldChar3)
    p._element.append(instrText2)
    p._element.append(fldChar4)

def main():
    """Función principal"""
    print("="*60)
    print("GENERANDO DOCUMENTO WORD PROFESIONAL")
    print("="*60)
    
    # Crear documento
    doc = Document()
    
    # Configurar márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)
    
    print("✓ Márgenes configurados")
    
    # Crear estilos
    crear_estilos(doc)
    print("✓ Estilos personalizados creados")
    
    # Agregar portada
    agregar_portada(doc)
    print("✓ Portada agregada")
    
    # Leer contenido del informe Markdown
    try:
        with open('VII. Documentos Transversales/INFORME_JUSTIFICACION_REDUCCION_CANTIDADES_v1.0.md', 'r', encoding='utf-8') as f:
            contenido_md = f.read()
        print("✓ Contenido Markdown leído")
    except Exception as e:
        print(f"❌ Error leyendo archivo: {e}")
        return
    
    # Procesar contenido
    print("🔄 Procesando contenido...")
    procesar_markdown_a_word(doc, contenido_md)
    print("✓ Contenido procesado")
    
    # Agregar numeración de páginas
    agregar_numeracion_paginas(doc)
    print("✓ Numeración de páginas agregada")
    
    # Guardar documento
    output_path = 'VII. Documentos Transversales/INFORME_JUSTIFICACION_REDUCCION_CANTIDADES_v1.0.docx'
    doc.save(output_path)
    print(f"\n✅ Documento guardado exitosamente en:")
    print(f"   {output_path}")
    print("\n" + "="*60)
    print("PROCESO COMPLETADO")
    print("="*60)

if __name__ == '__main__':
    main()






